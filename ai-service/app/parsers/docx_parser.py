"""Word Document Parser using python-docx to extract text sections and inline images."""

import io
import base64
from typing import List
import docx
from app.models import ExtractedPage, ExtractionResult


def parse_docx(file_path_or_bytes: bytes | str) -> ExtractionResult:
    """Extract text paragraphs, headings, tables, and images from a DOCX file.
    
    Groups content logically into sections/pages.
    
    Args:
        file_path_or_bytes: File path or raw bytes of the DOCX file.
        
    Returns:
        ExtractionResult containing section-by-section text and base64 images.
    """
    if isinstance(file_path_or_bytes, bytes):
        doc = docx.Document(io.BytesIO(file_path_or_bytes))
    else:
        doc = docx.Document(file_path_or_bytes)

    # Extract inline images from document relationships
    doc_images = []
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                image_bytes = rel.target_part.blob
                ext = rel.target_ref.split(".")[-1]
                if ext.lower() not in ["png", "jpg", "jpeg", "gif", "svg"]:
                    ext = "png"
                b64_str = base64.b64encode(image_bytes).decode("utf-8")
                data_uri = f"data:image/{ext};base64,{b64_str}"
                doc_images.append(data_uri)
            except Exception:
                pass

    # Group paragraphs by headings into logical sections
    sections: List[ExtractedPage] = []
    current_section_text = []
    section_number = 1

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Check if heading (starts a new logical section)
        if para.style.name.startswith("Heading") and current_section_text:
            sections.append(
                ExtractedPage(
                    page_number=section_number,
                    text="\n".join(current_section_text),
                    images=doc_images if section_number == 1 else []
                )
            )
            section_number += 1
            current_section_text = [text]
        else:
            current_section_text.append(text)

    # Add tables text if present
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                table_text.append(" | ".join(row_cells))
        if table_text:
            current_section_text.append("\n" + "\n".join(table_text))

    if current_section_text:
        sections.append(
            ExtractedPage(
                page_number=section_number,
                text="\n".join(current_section_text),
                images=doc_images if section_number == 1 else []
            )
        )

    # If no heading sections were created, create one fallback section
    if not sections:
        sections.append(
            ExtractedPage(
                page_number=1,
                text="Document text empty or unformatted.",
                images=doc_images
            )
        )

    return ExtractionResult(
        file_type="docx",
        total_pages=len(sections),
        pages=sections,
        metadata={"total_sections": len(sections)}
    )
